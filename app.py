import streamlit as st
import json
import random
import traceback
import re
import io
from docx import Document

class NeuroExpertMaster:
    def __init__(self, matrix_data):
        try:
            # СТЫКОВКА: Берем данные напрямую из памяти Стримлита
            self.lib = matrix_data
            
            def deep_find(data, target):
                if isinstance(data, dict):
                    if target in data: return data[target]
                    for v in data.values():
                        found = deep_find(v, target)
                        if found: return found
                return None
            
            self.rv = deep_find(self.lib, "risk_verification")
            self.sr = deep_find(self.lib, "suicide_risk")
            self.nv = deep_find(self.lib, "neuro_vectors")
            # st.toast — это маленькое всплывающее окно в Стримлите
            st.toast("✅ ПРИЛОЖЕНИЕ ЗАПУЩЕНО")
        except Exception as e: 
            st.error(f"❌ Ошибка инициализации: {e}")
            self.lib = {}

    def apply_gender(self, text, gender, is_endo=False):
        if isinstance(text, list):
            forbidden = "(органика)" if is_endo else "(эндоген)"
            # 1. Сначала ВЫКИДЫВАЕМ всё чуждое:
            filtered = [str(t) for t in text if forbidden not in str(t).lower()]
            # 2. А теперь выбираем только из того, что осталось:
            text = random.choice(filtered) if filtered else random.choice(text)

        text = str(text)
        is_fem = (gender == 'а')
        if is_fem: text = text.replace("ен{g}", "на").replace("{g}", "а")
        else: text = text.replace("{g}", "")

        fem_map = {"инертен": "инертна", "активен": "активна", "спокоен": "спокойна", "ориентирован": "ориентирована",
                   "адекватен": "адекватна", "неустойчив": "неустойчива", "заторможен": "заторможена",
                   "подвижен": "подвижна", "зависим": "зависима", "резв": "резва", "доступен": "доступна"}
        if is_fem:
            for m, f in fem_map.items(): text = re.sub(rf"\b{m}\b", f, text)
            text = text.replace("пациент ", "пациентка ")
        else: text = text.replace("пациентка", "пациент")

        text = text.replace("['", "").replace("']", "").replace("'", "").replace('"', "")
        text = re.sub(r"\(.*?\)", "", text).replace("..", ".").replace(" ,", ",").replace(",,", ",").replace(". ,", ". ")
        sentences = [s.strip().capitalize() for s in text.split('.') if s.strip()]
        return ". ".join(sentences) + "." if sentences else ""

    def run(self, code_str, pr_in, t_in):
        try:
            head, s_raw = code_str.split('/')
            raw_typ = head.rstrip('мж'); gen = 'а' if head.endswith('ж') else ''
            s = [int(x) for x in s_raw if x.isdigit()][:10]
            while len(s) < 10: s.append(0)

            st_log = self.lib.get("status_logic", {}); cl_bases = self.lib.get("clinical_bases", {})
            concl = cl_bases.get("conclusions", self.lib.get("conclusions", {}))
            adj_lib = self.lib.get("phenomenology_adjustments", {})
            intros = st_log.get("intros", {})

            typ_key = raw_typ if raw_typ in intros else ('0' if raw_typ.startswith('0') else raw_typ)
            is_endo = (typ_key == "8"); is_organ = typ_key in ["1","2","3","4","5"]
            is_norm_logic = typ_key.startswith('0') and sum(s) < 10
            presets = [p.strip() for p in pr_in.split(',') if p.strip()]; tags = [t.strip().lower() for t in t_in.split(',') if t.strip()]

            # --- 1. Психологический статус ---
            st_raw = [intros.get(typ_key, intros.get("0", "")), cl_bases.get(raw_typ if raw_typ in cl_bases else str(round(sum(s) / 10)), "")]
            # Добавляем ПА в статус, если есть тег
            if "па" in tags: st_raw.append(self.lib.get("tags", {}).get("па", ""))
            for p in presets:
                # Ищем надстройку и как ввел (ретик), и капсом (РЕТИК) на всякий случай
                p_data = adj_lib.get(p, adj_lib.get(p.upper(), {}))
                st_raw.append(p_data.get("status", ""))
            for t in tags:
                if t != "па": st_raw.append(self.lib.get("tags", {}).get(t, ""))
            if hasattr(self, 'sr') and self.sr:
                # Берем фразу по типу профиля (3, 8, 0*) или дефолт "0"
                s_phrase = self.sr.get(typ_key, self.sr.get("0", ""))
                if s_phrase:
                    st_raw.append(s_phrase)
            status_text = ". ".join([self.apply_gender(p, gen, is_endo).rstrip('.') for p in st_raw if p]) + "."

            # --- 2. РЕЗУЛЬТАТЫ ---
            f_res = []
            f_keys = ["attention", "visual_gnosis", "spatial", "dynamic_praxis", "afferent_praxis", "cube", "calculation", "speech", "memory", "thinking"]
            for i, k in enumerate(f_keys):
                f_res.append(self.apply_gender(self.lib.get("functions", {}).get(k, {}).get(str(s[i]), [""]), gen, is_endo))
                for p in presets:
                    res_adj = adj_lib.get(p, {}).get("results", "")
                    if res_adj and ((i == 3 and p in ["Апрдин", "леврег", "праврег"]) or (i == 7 and p in ["Асенс", "Асем", "Ааф"])):
                        f_res.append(self.apply_gender(res_adj, gen, is_endo))

            # --- 3. ЗАКЛЮЧЕНИЕ ---
            final = []
            # Вектор Депрессии/Горя
            if (typ_key == "9" or any(p.startswith('Д') for p in presets)) and self.nv:
                final.append(self.apply_gender(self.nv.get("9", ""), gen, is_endo))

            # САММАРИ (Синтез по баллам)
            summ_list = concl.get("functional_summaries", [""])
            s_sum = sum(s); s_idx = 0 if s_sum <= 7 else (1 if s_sum <= 15 else 2)
            if is_norm_logic: s_idx = 0
            final.append(self.apply_gender(summ_list[s_idx if s_idx < len(summ_list) else -1], gen, is_endo))

            # --- СИНТЕЗ (ЦЕМЕНТ) ---
            # 1. Берем данные. Если это словарь (старый вид) - превращаем в список значений
            raw_synth = concl.get("synthesis", [])
            synth_list = list(raw_synth.values()) if isinstance(raw_synth, dict) else raw_synth

            if synth_list and s_sum >= 3:
                # Определяем корзину (0, 1 или 2)
                s_idx = 1 if is_endo else (0 if s_sum <= 7 else (1 if s_sum <= 15 else 2))
                # Безопасный забор фразы (чтобы не было IndexError)
                target_phrase = synth_list[s_idx if s_idx < len(synth_list) else -1]
                final.append(self.apply_gender(target_phrase, gen, is_endo))

            # --- ВЕКТОРЫ (Аддитивная сборка: Личина + ВСЕ Теги + ПА) ---
            v_parts = []

            # 1. Базовый вектор (ТОЛЬКО для избранных 8, 7, 9)
            if is_endo or (typ_key in ["9", "7"]):
                base_v = self.nv.get(typ_key, "")
                if base_v: v_parts.append(base_v)

            # 2. А ТЕГИ ДОСТУПНЫ ВСЕМ (Выровнял влево!)
            for t in tags:
                if t in self.nv and t != "па":
                    v_parts.append(self.nv[t])

            # 3. ПА ТОЖЕ ДЛЯ ВСЕХ (Выровнял влево!)
            if "па" in tags and "па" in self.nv:
                v_parts.append(self.nv["па"])

            # 4. СБОРКА И ЧИСТКА (Выровнял влево!)
            if v_parts:
                v_clean = [random.choice(p) if isinstance(p, list) else p for p in v_parts]
                final.append(self.apply_gender(" ".join(v_clean), gen, is_endo))

            # ФАКТОРЫ С БУСТЕРАМИ
            if not is_norm_logic:
                f_active = []; f_p = concl.get("factors", {})
                b1 = 3 if any(p in ["н", "Апат", "асте"] for p in presets) and is_organ else 0
                b2 = 3 if any(p in ["Асенс", "Ааф", "Аак", "Асем", "Апркин", "Апркон", "АгнП", "АгнЛ", "неглект"] for p in presets) and is_organ else 0
                b3 = 3 if any(p in ["праврег", "леврег", "Аэф", "Апрдин"] for p in presets) and is_organ else 0

                if (s[0] + s[6] + b1 >= 3): f_active.append(self.apply_gender(f_p.get('neurodynamic', ''), gen, is_endo).rstrip('.') + " (I блок)")
                if (s[1] + s[2] + s[5] + b2 >= 3): f_active.append(self.apply_gender(f_p.get('spatial', ''), gen, is_endo).rstrip('.') + " (II блок)")
                if (s[3] + s[9] + b3 >= 3): f_active.append(self.apply_gender(f_p.get('regulatory', ''), gen, is_endo).rstrip('.') + " (III блок)")

                if f_active:
                    res_f = f_active[0]
                    for n_f in f_active[1:]: res_f = res_f.rstrip(".") + ", " + n_f[:1].lower() + n_f[1:]
                    for p in presets:
                        concl_adj = adj_lib.get(p, {}).get("conclusion", "")
                        if concl_adj: res_f = res_f.rstrip(".") + ". " + self.apply_gender(concl_adj, gen, is_endo)
                    final.append(res_f)

            # РЕЗЮМЕ И ВАРИАТИВНЫЙ РИСК
            final.append(self.apply_gender(concl.get("personality_emotional_resume", {}).get("stable" if is_norm_logic else "unstable", ""), gen, is_endo))
            if self.rv:
                risk_opts = self.rv.get(typ_key, self.rv.get("0", ""))
                final.append(self.apply_gender(risk_opts, gen, is_endo))

            return f"ПСИХОЛОГИЧЕСКИЙ СТАТУС:\n{status_text}\n\nРЕЗУЛЬТАТЫ:\n{' '.join(f_res)}\n\nЗАКЛЮЧЕНИЕ:\n" + "\n".join([p.strip() for p in final if p.strip()])
            return final_res

        except Exception: 
            return traceback.format_exc()

# Загрузка JSON-аккумулятора
@st.cache_data
def load_matrix():
    with open('expert_matrix.json', 'r', encoding='utf-8-sig') as f:
        return json.load(f)

matrix = load_matrix()

st.set_page_config(page_title="NeuroExpert Web", layout="wide")

import base64

# 1. ФУНКЦИЯ КОДИРОВАНИЯ КАРТИНКИ (Чтобы вшить в HTML)
def get_base64_image(image_path):
    try:
        with open(image_path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode()
    except:
        return ""

# Кодируем твой мозг
img_base64 = get_base64_image("brain3.jpg")

# 2. ЕДИНЫЙ БЛОК: ГРАДИЕНТ + МОЗГ + ТЕКСТ (ВСЁ В ОДНОМ СТАКАНЕ)
st.markdown(f"""
    <div style="
        background: linear-gradient(90deg, #0e1117 0%, #1c1f26 100%); 
        padding: 20px; 
        border-radius: 15px; 
        border-left: 5px solid #FF4B4B; 
        display: flex; 
        align-items: center; 
        gap: 20px;
        margin-bottom: 25px;
    ">
        <img src="data:image/jpeg;base64,{img_base64}" 
             style="width: 60px; height: 60px; border-radius: 10px; object-fit: cover;">
        <div>
            <h1 style="color: #ffffff; margin: 0; font-family: 'Segoe UI'; font-size: 2.2em; line-height: 1.1;">
                <span style="color: #FF4B4B;">Neuro</span>Expert
            </h1>
            <p style="color: #808495; font-style: italic; margin-top: 4px; font-size: 0.95em;">
                Комплексная система синдромального нейропсихологического анализа
            </p>
        </div>
    </div>
""", unsafe_allow_html=True)

# 3. ГАЙД - В САЙДБАР (ЧТОБЫ НЕ ПОРТИЛ МОНОЛИТ)
with st.sidebar:
    st.markdown("---")
    try:
        with open("AppGuide.pdf", "rb") as f:
            st.download_button("📚 СКАЧАТЬ ГАЙД (PDF)", f, "NeuroExpert_Guide.pdf", "application/pdf")
    except:
        pass

# --- 2. ЛЕВАЯ ПАНЕЛЬ (САЙДБАР) ---
with st.sidebar:
    # Твои колонки для мозгов
    c1, c2, c3 = st.columns([1, 2, 1]) 
    with c2:
        try:
            st.image("brain2.jpg", width=150)
        except:
            st.write("🧠")
            
    st.header("📋 Паспорт")
    
    # Кнопка сброса (одна, в сайдбаре)
    if st.button("♻️ СБРОСИТЬ ВСЁ", type="secondary"):
        reset_app()
    
    fio = st.text_input("ФИО", "Иванов И.И.", key="fio_input")
    age = st.number_input("Возраст", 1, 110, 65, key="age_input")
    p_type = st.selectbox("Тип", ["0*", "0+", "00", "0т", "0-", "0000", "0", "0сон", "1", "2", "3", "4", "5", "7", "8", "9", "9гэ"], key="profile_select")
    p_gen = st.radio("Пол", ["м", "ж"], horizontal=True)

    st.markdown("---") # Черта
    
    # ПЕРЕНЕСЛИ СПИСКИ СЮДА (ДЛЯ УДОБСТВА И СМАРТФОНА)
    adj_keys = list(matrix.get("phenomenology_adjustments", {}).keys())
    presets = st.multiselect("🛠 Надстройки", adj_keys, key="adj_ms")
    tag_keys = list(matrix.get("tags", {}).keys())
    selected_tags = st.multiselect("🏷 Теги", tag_keys, key="tags_ms")

# --- 3. ЦЕНТРАЛЬНОЕ ПОЛЕ (ФУНКЦИИ) ---
st.subheader("📊 Функции (0-5)")
f_names = ["Внимание", "Зрит.Гнозис", "Простр.Гнозис", "Дин.Праксис", "Кин.Праксис", "Констр.Праксис", "Счет", "Речь", "Память", "Мышление"]
scores = []

# УБРАЛИ КОЛОНКИ: теперь ползунки идут СТРОГО 1, 2, 3... без путаницы на смарте
for i, name in enumerate(f_names):
    scores.append(st.slider(f"{i+1}. {name}", 0, 5, 0, key=f"s_{i}"))

# --- 4. КНОПКА ГЕНЕРАЦИИ (В САМОМ НИЗУ) ---
if st.button("🚀 СГЕНЕРИРОВАТЬ ПРОТОКОЛ"):
    # Твой движок запускается здесь...
    full_code = f"{p_type}{p_gen}/{''.join(map(str, scores))}"
    
    # Создаем объект энджина прямо здесь
    engine = NeuroExpertMaster(matrix)
    
    # Прогоняем через RUN
    report = engine.run(full_code, ",".join(presets), ",".join(selected_tags))
    
    st.markdown("### Итоговый текст:")
    st.text_area("", report, height=500)
    
    # ВОРД-ПРИНТЕР
    doc = Document()
    doc.add_paragraph(f"ПРОТОКОЛ: {fio}")
    doc.add_paragraph(report)
    bio = io.BytesIO()
    doc.save(bio)
    st.download_button("📥 Скачать .docx", bio.getvalue(), f"{fio}.docx")
